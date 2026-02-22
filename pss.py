import streamlit as st
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
import pytesseract
import cv2
import numpy as np

st.set_page_config(page_title="Employee Onboarding", layout="wide")

# ---------- Premium Theme ----------
st.markdown("""
<style>
.main {
    background-color: #f5f7fb;
}
h1, h2, h3 {
    color: #2b2b2b;
}
.stButton>button {
    background-color: #635BFF;
    color: white;
    border-radius: 8px;
    height: 3em;
    width: 100%;
    font-size: 18px;
}
</style>
""", unsafe_allow_html=True)

st.title("🚀 Employee Joining Form (Premium Hinglish)")

# ---------- Stepper ----------
step = st.session_state.get("step", 1)
progress = int((step/5)*100)
st.progress(progress)
st.write(f"Step {step} / 5")

# ---------- STEP 1 ----------
if step == 1:
    st.header("👤 Basic Details")
    name = st.text_input("Poora Naam")
    mobile = st.text_input("Mobile Number")
    email = st.text_input("Email ID")
    dob = st.date_input("Date of Birth")
    father = st.text_input("Father Name")
    address = st.text_area("Current Address")
    p_address = st.text_area("Permanent Address")
    marital = st.selectbox("Marital Status", ["Single", "Married"])
    blood = st.text_input("Blood Group")

    selfie = st.camera_input("📸 Apni Selfie Le")

    if st.button("Next ➡️"):
        st.session_state.update(locals())
        st.session_state.step = 2
        st.rerun()

# ---------- STEP 2 Aadhaar ----------
elif step == 2:
    st.header("📇 Aadhaar Capture + OCR")

    aadhaar_front = st.camera_input("Aadhaar Front Photo")
    aadhaar_back = st.camera_input("Aadhaar Back Photo")

    aadhaar_number = ""

    if aadhaar_front:
        file_bytes = np.asarray(bytearray(aadhaar_front.read()), dtype=np.uint8)
        img = cv2.imdecode(file_bytes, 1)
        text = pytesseract.image_to_string(img)
        st.write("🔍 OCR Result:", text)

        import re
        match = re.search(r"\d{4}\s?\d{4}\s?\d{4}", text)
        if match:
            aadhaar_number = match.group()
            st.success(f"Aadhaar Auto Detected: {aadhaar_number}")

    manual_aadhaar = st.text_input("Confirm Aadhaar Number", value=aadhaar_number)

    if st.button("Next ➡️"):
        st.session_state.update(locals())
        st.session_state.step = 3
        st.rerun()

# ---------- STEP 3 PAN ----------
elif step == 3:
    st.header("🪪 PAN Capture")

    pan_front = st.camera_input("PAN Front Photo")
    pan_back = st.camera_input("PAN Back Photo")
    pan = st.text_input("PAN Number")

    if st.button("Next ➡️"):
        st.session_state.update(locals())
        st.session_state.step = 4
        st.rerun()

# ---------- STEP 4 Documents + Signature ----------
elif step == 4:
    st.header("📂 Other Documents + Signature")

    resume = st.file_uploader("Resume / CV")
    salary_slip = st.file_uploader("Last Salary Slip")
    tenth = st.file_uploader("10th Certificate")
    twelfth = st.file_uploader("12th Certificate")
    grad = st.file_uploader("Graduation Degree")
    master = st.file_uploader("Master Degree")
    school_leave = st.file_uploader("School Leaving Certificate")

    signature = st.camera_input("✍ Digital Signature Capture (white paper pe sign kare)")

    if st.button("Next ➡️"):
        st.session_state.update(locals())
        st.session_state.step = 5
        st.rerun()

# ---------- STEP 5 Office + DOCX ----------
elif step == 5:
    st.header("🏢 Office Use Only")

    emp_id = st.text_input("Employee ID")
    dept = st.text_input("Department")
    designation = st.text_input("Designation")
    doj = st.date_input("Date of Joining")
    salary = st.number_input("Salary Offered")
    bank = st.text_input("Bank Name")
    acc = st.text_input("Account Number")
    ifsc = st.text_input("IFSC Code")
    remarks = st.text_area("Remarks")

    def generate_docx():
        doc = Document()
        doc.add_heading("Employee Joining Form", level=1)

        # Employee Table
        doc.add_heading("Employee Details", level=2)
        table = doc.add_table(rows=0, cols=2)

        def add_row(label, value):
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)

        s = st.session_state
        add_row("Full Name", s.name)
        add_row("Mobile", s.mobile)
        add_row("Email", s.email)
        add_row("DOB", s.dob)
        add_row("Father Name", s.father)
        add_row("Address", s.address)
        add_row("Permanent Address", s.p_address)
        add_row("Marital", s.marital)
        add_row("Blood Group", s.blood)
        add_row("Aadhaar", s.manual_aadhaar)
        add_row("PAN", s.pan)

        # Images
        def add_image(title, img_file):
            if img_file:
                doc.add_heading(title, level=3)
                image = Image.open(img_file)
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format='PNG')
                doc.add_picture(img_byte_arr, width=Inches(2))

        add_image("Selfie", s.selfie)
        add_image("Aadhaar Front", s.aadhaar_front)
        add_image("Aadhaar Back", s.aadhaar_back)
        add_image("PAN Front", s.pan_front)
        add_image("PAN Back", s.pan_back)
        add_image("Signature", s.signature)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    if st.button("📄 Generate Editable DOCX"):
        file = generate_docx()
        st.download_button(
            "⬇️ Download DOCX",
            file,
            file_name=f"{st.session_state.name}_Joining_Form.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

import streamlit as st
from docx import Document
import os
import mammoth
import io
from datetime import datetime
import calendar
import random
import pandas as pd

# ---------------------------
# Helper: placeholder replacement preserving style
# ---------------------------
def advanced_replace_text_preserving_style(doc, replacements):
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        for key, value in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                new_runs = []
                accumulated = ""
                for run in runs:
                    accumulated += run.text
                    new_runs.append(run)
                    if placeholder in accumulated:
                        style_run = next((r for r in new_runs if placeholder in r.text), new_runs[0])
                        font = style_run.font
                        accumulated = accumulated.replace(placeholder, value)
                        for r in new_runs:
                            r.text = ''
                        if new_runs:
                            new_run = new_runs[0]
                            new_run.text = accumulated
                            try:
                                new_run.font.name = font.name
                                new_run.font.size = font.size
                                new_run.font.bold = font.bold
                                new_run.font.italic = font.italic
                                new_run.font.underline = font.underline
                                new_run.font.color.rgb = font.color.rgb
                            except:
                                pass
                        break

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

# ---------------------------
# Best Before Calculator (ADDED)
# ---------------------------
def calculate_best_before(date_str):
    try:
        date_str = date_str.strip().upper()
        parts = date_str.split()

        if len(parts) != 2:
            return ""

        month_name, year = parts
        year = int(year)

        month_num = list(calendar.month_name).index(month_name.capitalize())

        # Add 23 months (2 years - 1 month)
        new_month = month_num + 23
        new_year = year + (new_month - 1) // 12
        new_month = ((new_month - 1) % 12) + 1

        best_before = f"{calendar.month_name[new_month].upper()} {new_year}"
        return best_before

    except:
        return ""

# ---------------------------
# DOCX generate / preview
# ---------------------------
def generate_docx(data, template_path="template.docx", output_path="generated_coa.docx"):
    doc = Document(template_path)
    advanced_replace_text_preserving_style(doc, data)
    doc.save(output_path)
    return output_path

def docx_to_html(docx_path):
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        return result.value

# ---------------------------
# Distribution helper
# ---------------------------
def distribute_within_bounds(target, names, mins, maxs, weights):
    eps = 1e-9
    min_sum = sum(mins[n] for n in names)
    max_sum = sum(maxs[n] for n in names)
    if target + eps < min_sum or target - eps > max_sum:
        raise ValueError("Target not achievable with given bounds.")

    w_sum = sum(weights[n] for n in names)
    vals = {n: target * (weights[n] / w_sum) for n in names} if w_sum > 0 else {n: target / len(names) for n in names}

    locked = {n: False for n in names}
    for _ in range(100):
        changed = False
        for n in names:
            if not locked[n]:
                if vals[n] < mins[n]:
                    vals[n] = mins[n]
                    locked[n] = True
                    changed = True
                elif vals[n] > maxs[n]:
                    vals[n] = maxs[n]
                    locked[n] = True
                    changed = True
        if not changed:
            unlocked = [n for n in names if not locked[n]]
            if not unlocked:
                break
            rem = target - sum(vals.values())
            if abs(rem) < 1e-8:
                break
            w_un_sum = sum(weights[n] for n in unlocked)
            for n in unlocked:
                vals[n] += rem * (weights[n] / w_un_sum) if w_un_sum > 0 else rem / len(unlocked)

    for n in names:
        vals[n] = round(max(min(vals[n], maxs[n]), mins[n]), 2)

    if abs(sum(vals.values()) - target) > 0.01:
        raise ValueError("Could not distribute to meet target.")
    return vals

# ---------------------------
# Ranges
# ---------------------------
RANGES = {
    "fat": (0.45, 0.55),
    "air": (2.90, 3.10),
    "ash": (0.45, 0.55),
    "protein": (2.45, 2.55),
    "gum": (80.10, 89.95)
}

# ---------------------------
# Random calculation
# ---------------------------
def calculate_components_random(moisture):
    remaining = 100 - moisture
    others = ["fat", "air", "ash", "protein"]
    for _ in range(2000):
        sampled = {o: round(random.uniform(*RANGES[o]), 2) for o in others}
        gum = round(remaining - sum(sampled.values()), 2)
        if RANGES["gum"][0] <= gum <= RANGES["gum"][1]:
            return gum, sampled["protein"], sampled["ash"], sampled["air"], sampled["fat"]
    raise ValueError("Failed to calculate components")

# ---------------------------
# Streamlit UI
# ---------------------------
st.set_page_config(page_title="COA Generator", layout="wide")

code = st.selectbox("QUALITY choose karo", [f"{i}-{i+500}" for i in range(500, 10001, 500)])
date = st.text_input("DATE (JULY 2025)")
batch_no = st.text_input("BATCH NUMBER (A/25/11011)")
moisture = st.number_input("MOISTURE (%)", 0.0, 99.0, 10.0, step=0.01)

ph = st.text_input("pH")
mesh_200 = st.text_input("200# MESH (%)")
viscosity_2h = st.text_input("VISCOSITY 2HRS")
viscosity_24h = st.text_input("VISCOSITY 24HRS")

if st.button("Generate COA"):
    gum, protein, ash, air, fat = calculate_components_random(moisture)

    # ✅ AUTO BEST BEFORE
    best_before = calculate_best_before(date)

    data = {
        "DATE": date,
        "BEST_BEFORE": best_before,  # 👈 added
        "BATCH_NO": batch_no,
        "MOISTURE": f"{moisture:.2f}%",
        "PH": ph,
        "MESH_200": f"{mesh_200}%",
        "VISCOSITY_2H": viscosity_2h,
        "VISCOSITY_24H": viscosity_24h,
        "GUM_CONTENT": f"{gum:.2f}%",
        "PROTEIN": f"{protein:.2f}%",
        "ASH_CONTENT": f"{ash:.2f}%",
        "AIR": f"{air:.2f}%",
        "FAT": f"{fat:.2f}%"
    }

    template_path = f"COA {code}.docx"
    output_path = "generated_coa.docx"

    if os.path.exists(template_path):
        generate_docx(data, template_path, output_path)

        safe_batch = (batch_no or "BATCH").replace("/", "-").replace("\\", "-")
        final_filename = f"COA {safe_batch} {code}.docx"

        with open(output_path, "rb") as f:
            st.download_button(
                "📥 Download COA",
                f.read(),
                file_name=final_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("Template not found")
